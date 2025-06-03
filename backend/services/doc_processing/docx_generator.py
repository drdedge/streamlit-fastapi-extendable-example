from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from ...models.doc_extraction import ProcessingMetadata

class DocxGenerator:
    def create_document(self, content: str, metadata: ProcessingMetadata) -> Document:
        doc = Document()
        
        # Add title
        title = doc.add_heading('Extracted Document', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata section
        doc.add_heading('Document Information', level=1)
        
        # Create metadata table
        table = doc.add_table(rows=7, cols=2)
        table.style = 'Light List Accent 1'
        
        # Fill metadata
        metadata_items = [
            ('Original File', metadata.filename),
            ('File Size', f"{metadata.file_size:,} bytes"),
            ('Upload Time', metadata.upload_datetime.strftime('%Y-%m-%d %H:%M:%S')),
            ('Processing Time', metadata.processing_datetime.strftime('%Y-%m-%d %H:%M:%S')),
            ('Issuer ID', metadata.issuer_id),
            ('User Email', metadata.user_email),
            ('Original Format', metadata.original_format.upper())
        ]
        
        for idx, (label, value) in enumerate(metadata_items):
            table.cell(idx, 0).text = label
            table.cell(idx, 1).text = str(value)
        
        # Add content section
        doc.add_page_break()
        doc.add_heading('Document Content', level=1)
        
        # Add the actual content
        content_paragraph = doc.add_paragraph(content)
        content_paragraph.style = 'Normal'
        
        # Add footer
        doc.add_page_break()
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.add_run(f'Generated on {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}').italic = True
        
        return doc
    
    def save_document(self, doc: Document, output_path: str) -> str:
        doc.save(output_path)
        return output_path